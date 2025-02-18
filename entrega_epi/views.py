from django.shortcuts import get_object_or_404, redirect
from django.conf import settings
from django.http import JsonResponse
from django.contrib import messages
from django.core.files.base import ContentFile
from django.views.generic import ListView, CreateView,View
from django.utils.encoding import smart_str
import os
import requests
from io import BytesIO
from datetime import datetime
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from django.http import HttpResponse
from entrega_epi.models import EntregaColaboradorEPI, FichaEPI
from colaboradores.models import Colaborador
from entrega_epi.models import FichaEPI
from django.core.files.storage import default_storage

# Create your views here.

class MostrarListaColaboradoresView(ListView):
    model = Colaborador
    template_name = 'entrega_epi/entrega_lista.html'
    context_object_name = 'colaboradores'

    def get_queryset(self):
        queryset = super().get_queryset()
        # Adiciona os atributos `tem_epis` a cada colaborador
        for colaborador in queryset:
            colaborador.tem_epis = colaborador.entregas.exists()  # Verifica se há entregas para o colaborador
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        return context

class AdicionarEPIView(CreateView):
    model = EntregaColaboradorEPI
    fields = ['epi',]
    template_name = 'entrega_epi/adicionar_epi.html'

    def form_valid(self, form):
        colaborador_slug = self.kwargs['slug']
        colaborador = get_object_or_404(Colaborador, slug=colaborador_slug)
        epi = form.cleaned_data['epi']  

        if EntregaColaboradorEPI.objects.filter(colaborador=colaborador, epi=epi).exists():
            return super().form_invalid(form) 

        form.instance.colaborador = colaborador  
        form.save()
        return self.render_to_response(self.get_context_data(form=form))

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        colaborador_slug = self.kwargs['slug']
        colaborador = get_object_or_404(Colaborador, slug=colaborador_slug)
        context['colaborador'] = colaborador
        entregas = EntregaColaboradorEPI.objects.filter(colaborador=colaborador)
        context['entregas'] = entregas
        ficha_recente = FichaEPI.objects.filter(colaborador=colaborador).order_by('-data_geracao').first()
        context['ficha_recente'] = ficha_recente
        
        return context

class LimparEntregasView(View):
    def post(self, request, slug, *args, **kwargs):
        colaborador = get_object_or_404(Colaborador, slug=slug)
        EntregaColaboradorEPI.objects.filter(colaborador=colaborador).delete()
        return redirect('adicionar_epi', slug=slug)


class GerarFormularioEPIView(View):
    def get(self, request, slug, *args, **kwargs):
        # Validar o colaborador
        colaborador = get_object_or_404(Colaborador, slug=slug)
        entregas = EntregaColaboradorEPI.objects.filter(colaborador=colaborador)
        epis = [entrega.epi for entrega in entregas]

        # Caminho para o modelo do documento
        modelo_url = default_storage.url("documentos/Ficha de Entrega de EPIs - formulário.docx")
        
        response = requests.head(modelo_url)

        if response.status_code != 200:
            raise FileNotFoundError(f"Arquivo do modelo não encontrado: {modelo_url}")

        # Fazer o download do arquivo
        response = requests.get(modelo_url)
        if response.status_code == 200:
            temp_path = "/tmp/modelo.docx"  # Caminho temporário no servidor Render
            with open(temp_path, "wb") as f:
                f.write(response.content)

            modelo_path = temp_path  # Agora o arquivo pode ser acessado localmente
        else:
            raise FileNotFoundError(f"Erro ao baixar o modelo do GCS: {modelo_url}")

        # Carregar o documento Word
        doc = Document(modelo_path)

        # Substituir os campos de texto na tabela
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "Nome completo:" in cell.text:
                        cell.text = f"Nome completo: {colaborador.nome}"
                    elif "Função:" in cell.text:
                        cell.text = f"Função: {colaborador.cargo}"

        # Localizar a tabela de EPIs
        table = None
        for tbl in doc.tables:
            if "Descrição Simplificada" in tbl.cell(0, 0).text:
                table = tbl
                break

        if table is None:
            raise ValueError("Tabela de EPIs não encontrada no modelo.")

        # Preencher a tabela com os EPIs entregues
        for epi in epis:
            row_cells = table.add_row().cells
            row_cells[0].text = epi.nome
            row_cells[1].text = "1"
            row_cells[2].text = epi.numero_ca
            row_cells[3].text = epi.modelo
            row_cells[4].text = epi.fabricante if epi.fabricante else "N/A"
            row_cells[5].text = epi.validade_ca.strftime("%d/%m/%Y") if epi.validade_ca else "N/A"
            row_cells[6].text = datetime.now().strftime("%d/%m/%Y")

        # Criar um buffer de memória para armazenar o arquivo .docx
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        # Nome do arquivo
        nome_arquivo = f'Ficha_EPI_{colaborador.nome}_{datetime.now().strftime("%Y%m%d%H%M%S")}.docx'

        # Salvar a ficha no banco de dados
        ficha = FichaEPI.objects.create(colaborador=colaborador)
        ficha.arquivo.save(nome_arquivo, ContentFile(doc_buffer.getvalue()))

        # Criar a resposta de download
        response = HttpResponse(doc_buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{smart_str(nome_arquivo)}"'
        doc_buffer.close()

        return response


class ApagarFichaEPIView(View):
    def post(self, request, ficha_id, *args, **kwargs):
        # Buscar a ficha pelo ID
        ficha = get_object_or_404(FichaEPI, id=ficha_id)

        
        if ficha.arquivo:
            ficha.arquivo.delete()

        # Apagar o registro do banco de dados
        ficha.delete()

        # Adicionar uma mensagem de sucesso
        messages.success(request, "Ficha de EPI apagada com sucesso!")

        # Redirecionar para a página de lista de entregas
        return redirect('entrega_lista')

from entrega_epi.models import FichaEPI

class FichaRecenteAPI(View):
    def get(self, request, colaborador_slug):
        colaborador = get_object_or_404(Colaborador, slug=colaborador_slug)
        ficha = FichaEPI.objects.filter(colaborador=colaborador).order_by('-data_geracao').first()
        if ficha:
            return JsonResponse({
                'id': ficha.id,
                'data_geracao': ficha.data_geracao.strftime('%Y-%m-%d'),
                'arquivo_url': ficha.arquivo.url
            })
        return JsonResponse({'message': 'Nenhuma ficha encontrada'}, status=404)

